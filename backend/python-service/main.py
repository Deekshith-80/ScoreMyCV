from __future__ import annotations

import io
import json
import logging
import os
import re
import textwrap
import urllib.error
import urllib.request
from collections import Counter
from typing import Any, Dict, List, Optional, Tuple

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel, Field
from PyPDF2 import PdfReader
from docx import Document
import spacy

try:
    from sentence_transformers import SentenceTransformer
except Exception:  # pragma: no cover - optional dependency
    SentenceTransformer = None


APP_NAME = "ATS Service"
FRONTEND_ORIGIN = os.getenv("FRONTEND_ORIGIN", "*")
ALLOWED_ORIGINS = [
    origin.strip()
    for origin in os.getenv(
        "ALLOWED_ORIGINS",
        "http://localhost:5173,https://resumepilot-ai-frontend.vercel.app,http://localhost:5000,https://resumepilot-backend-api.vercel.app",
    ).split(",")
    if origin.strip()
]
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1-mini").strip()
ATS_EXPORT_MAX_LINES = int(os.getenv("ATS_EXPORT_MAX_LINES", "55"))

logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s %(levelname)s %(name)s %(message)s",
)
logger = logging.getLogger("ats-service")

app = FastAPI(title=APP_NAME, version="2.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS if FRONTEND_ORIGIN == "*" else [FRONTEND_ORIGIN, *ALLOWED_ORIGINS],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

try:
    NLP = spacy.load("en_core_web_sm")
except Exception:
    NLP = spacy.blank("en")

try:
    EMBEDDER = SentenceTransformer("all-MiniLM-L6-v2") if SentenceTransformer else None
except Exception:
    EMBEDDER = None

ATS_SKILLS = [
    "python",
    "javascript",
    "typescript",
    "react",
    "node",
    "node.js",
    "express",
    "mongodb",
    "sql",
    "postgresql",
    "mysql",
    "aws",
    "azure",
    "gcp",
    "docker",
    "kubernetes",
    "fastapi",
    "flask",
    "django",
    "figma",
    "ux",
    "ui",
    "machine learning",
    "nlp",
    "data analysis",
    "testing",
    "automation",
    "cypress",
    "playwright",
    "ci/cd",
    "terraform",
    "html",
    "css",
    "tailwind",
    "git",
    "redis",
    "graphql",
    "rest api",
    "microservices",
    "agile",
    "scrum",
    "linux",
    "bash",
    "kafka",
    "elasticsearch",
    "spark",
    "pandas",
    "numpy",
]

CERTIFICATION_KEYWORDS = [
    "aws certified",
    "azure certified",
    "gcp certified",
    "pmp",
    "scrum",
    "csm",
    "ckad",
    "cka",
    "security+",
    "network+",
    "agile",
    "scrum master",
    "terraform associate",
    "google cloud",
    "microsoft certified",
]

SECTION_HEADINGS = [
    "summary",
    "profile",
    "experience",
    "work experience",
    "education",
    "projects",
    "skills",
    "certifications",
    "contact",
    "achievements",
]

JOB_REQUIREMENT_HINTS = [
    "required",
    "responsibilities",
    "qualifications",
    "must have",
    "you will",
    "looking for",
    "experience with",
    "years of experience",
]

ACTION_VERBS = [
    "built",
    "created",
    "developed",
    "designed",
    "launched",
    "implemented",
    "improved",
    "optimized",
    "led",
    "reduced",
    "increased",
    "delivered",
    "automated",
    "scaled",
]


class CoverLetterRequest(BaseModel):
    resume_summary: str = ""
    job_title: str = ""
    company_name: str = ""
    job_description: str = ""
    skills: List[str] = Field(default_factory=list)
    experience: List[str] = Field(default_factory=list)


class MatchJobRequest(BaseModel):
    resume_text: str = ""
    resume_skills: List[str] = Field(default_factory=list)
    resume_experience: List[str] = Field(default_factory=list)
    job_title: str = ""
    job_description: str = ""
    required_skills: List[str] = Field(default_factory=list)
    required_certifications: List[str] = Field(default_factory=list)
    location: str = ""
    experience: str = ""


class ExportRequest(BaseModel):
    content: str = ""
    file_name: str = "optimized-resume"
    format: str = "pdf"
    title: str = "Resume"
    kind: str = "resume"


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9+\-./\s]", " ", text.lower())).strip()


def split_lines(text: str) -> List[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def unique_preserve_order(items: List[str]) -> List[str]:
    seen = set()
    ordered = []
    for item in items:
        cleaned = item.strip()
        key = normalize(cleaned)
        if cleaned and key and key not in seen:
            seen.add(key)
            ordered.append(cleaned)
    return ordered


def chunked(items: List[str], size: int) -> List[List[str]]:
    return [items[index : index + size] for index in range(0, len(items), size)]


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception:
        logger.exception("Failed to parse PDF")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        return "\n".join(paragraphs)
    except Exception:
        logger.exception("Failed to parse DOCX")
        return ""


def extract_text(file: UploadFile, file_bytes: bytes) -> str:
    filename = (file.filename or "").lower()
    content_type = (file.content_type or "").lower()

    if filename.endswith(".pdf") or content_type == "application/pdf":
        text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx") or "wordprocessingml.document" in content_type:
        text = extract_text_from_docx(file_bytes)
    else:
        text = ""

    if not text.strip():
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = ""

    return text.replace("\x00", "").strip()


def extract_email(text: str) -> str:
    match = re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text, re.IGNORECASE)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    patterns = [
        r"(\+\d{1,3}[\s-]?)?(?:\(?\d{3}\)?[\s-]?)\d{3}[\s-]?\d{4}",
        r"\b\d{10}\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    return ""


def extract_name(text: str) -> str:
    lines = split_lines(text)
    if not lines:
        return ""

    for line in lines[:8]:
        cleaned = re.sub(r"[^A-Za-z\s'.-]", " ", line).strip()
        if not cleaned or extract_email(cleaned) or extract_phone(cleaned):
            continue
        tokens = [token for token in cleaned.split() if token]
        if 2 <= len(tokens) <= 4 and all(token[0].isupper() or token.isupper() for token in tokens):
            return cleaned
    return lines[0][:80]


def find_section_lines(text: str, section_names: List[str]) -> Dict[str, List[str]]:
    lines = split_lines(text)
    sections: Dict[str, List[str]] = {name: [] for name in section_names}
    current = None

    for line in lines:
        normalized = normalize(line).rstrip(":")
        matched = next((name for name in section_names if normalized == name or normalized.startswith(name + " ")), None)
        if matched:
            current = matched
            continue
        if current:
            sections[current].append(line)

    return sections


def extract_skills(text: str, job_description: str = "") -> List[str]:
    haystack = normalize(f"{text} {job_description}")
    matches: List[str] = []

    for skill in ATS_SKILLS:
        if re.search(rf"(^|[^a-z0-9]){re.escape(skill)}([^a-z0-9]|$)", haystack):
            matches.append(skill)

    sections = find_section_lines(text, ["skills"])
    for line in sections.get("skills", []):
        lower = normalize(line)
        for skill in ATS_SKILLS:
            if skill in lower:
                matches.append(skill)

    try:
        if NLP and getattr(NLP, "pipe_names", None):
            doc = NLP(text)
            for chunk in getattr(doc, "noun_chunks", []):
                chunk_text = chunk.text.strip()
                if 2 <= len(chunk_text) <= 40 and any(char.isalpha() for char in chunk_text):
                    matches.append(chunk_text)
    except Exception:
        logger.exception("Skill extraction via NLP failed")

    return unique_preserve_order(matches)


def extract_certifications(text: str) -> List[str]:
    lines = split_lines(text)
    sections = find_section_lines(text, ["certifications"])
    candidates = sections.get("certifications", [])
    for line in lines:
        lower = normalize(line)
        if any(keyword in lower for keyword in CERTIFICATION_KEYWORDS):
            candidates.append(line)
    return unique_preserve_order(candidates[:10])


def extract_education(text: str) -> List[str]:
    lines = split_lines(text)
    keywords = ["bachelor", "master", "degree", "b.sc", "m.sc", "phd", "diploma", "university", "college", "education"]
    education_lines = [line for line in lines if any(keyword in line.lower() for keyword in keywords)]
    return unique_preserve_order(education_lines[:8])


def extract_experience(text: str) -> List[str]:
    lines = split_lines(text)
    sections = find_section_lines(text, ["experience", "work experience", "professional experience"])
    candidates = sections.get("experience", []) + sections.get("work experience", []) + sections.get("professional experience", [])

    for line in lines:
        lower = normalize(line)
        if (
            any(verb in lower for verb in ACTION_VERBS)
            or re.search(r"\b(\d+)\+?\s*years?\b", lower)
            or re.search(r"\b20\d{2}\b", lower)
        ):
            candidates.append(line)

    return unique_preserve_order(candidates[:15])


def extract_projects(text: str) -> List[str]:
    lines = split_lines(text)
    sections = find_section_lines(text, ["projects"])
    candidates = sections.get("projects", [])
    for line in lines:
        lower = normalize(line)
        if "project" in lower or any(verb in lower for verb in ["built", "created", "developed", "launched", "shipped", "architected"]):
            candidates.append(line)
    return unique_preserve_order(candidates[:10])


def extract_job_requirements(job_description: str, required_skills: Optional[List[str]] = None) -> Tuple[List[str], List[str], List[str]]:
    normalized = normalize(job_description)
    required_skills = required_skills or []
    matched_skills = []
    required_keywords = []
    required_certs = []

    for skill in ATS_SKILLS:
        if skill in normalized:
            required_skills.append(skill)

    for item in unique_preserve_order(required_skills):
        if item:
            matched_skills.append(item)

    tokens = [token for token in re.split(r"[^a-z0-9+.#/-]+", normalized) if len(token) > 3]
    counts = Counter(tokens)
    for token, count in counts.items():
        if count >= 2 or token in {"docker", "kubernetes", "redis", "aws", "ci/cd", "graphql", "terraform", "testing", "fastapi", "react", "node", "mongodb"}:
            required_keywords.append(token)

    for phrase in CERTIFICATION_KEYWORDS:
        if phrase in normalized:
            required_certs.append(phrase)

    if not required_keywords and normalized:
        for line in split_lines(job_description):
            lower = normalize(line)
            if any(hint in lower for hint in JOB_REQUIREMENT_HINTS):
                required_keywords.extend([token for token in re.split(r"[^a-z0-9+.#/-]+", lower) if len(token) > 3])

    return unique_preserve_order(matched_skills), unique_preserve_order(required_keywords), unique_preserve_order(required_certs)


def keyword_score(resume_text: str, job_description: str, resume_skills: List[str], job_skills: List[str]) -> Tuple[int, List[str]]:
    if not job_description.strip():
        text_tokens = set(normalize(resume_text).split())
        skill_hits = len([skill for skill in ATS_SKILLS if skill in text_tokens])
        return max(30 if resume_text.strip() else 0, min(100, 35 + skill_hits * 3)), []

    resume_tokens = set(re.split(r"[^a-z0-9+.#/-]+", normalize(f"{resume_text} {' '.join(resume_skills)}")))
    job_tokens = [token for token in re.split(r"[^a-z0-9+.#/-]+", normalize(job_description)) if len(token) > 3]
    keyword_candidates = unique_preserve_order(job_skills + [token for token in job_tokens if token not in {"responsibilities", "qualifications", "experience", "required", "preferred"}])

    hits = []
    missing = []
    for keyword in keyword_candidates:
        normalized = normalize(keyword)
        if not normalized:
            continue
        if normalized in resume_tokens or normalized in normalize(resume_text):
            hits.append(keyword)
        else:
            missing.append(keyword)

    score = int(round(min(100, 15 + (len(hits) / max(1, len(keyword_candidates))) * 85)))
    return score, missing[:15]


def skills_score(resume_skills: List[str], job_skills: List[str]) -> Tuple[int, List[str]]:
    if not resume_skills and not job_skills:
        return 35, []

    resume_norm = {normalize(skill) for skill in resume_skills}
    job_norm = [normalize(skill) for skill in job_skills if normalize(skill)]

    if not job_norm:
        score = min(100, 40 + len(resume_skills) * 4)
        return score, []

    matching = [skill for skill in job_norm if skill in resume_norm]
    missing = [skill for skill in job_skills if normalize(skill) not in resume_norm]
    score = int(round(min(100, 10 + (len(matching) / max(1, len(job_norm))) * 90)))
    return score, missing[:12]


def experience_score(experience: List[str], text: str, job_description: str = "") -> Tuple[int, List[str]]:
    experience_items = experience or extract_experience(text)
    years_mentions = len(re.findall(r"\b\d+\+?\s*years?\b", text.lower()))
    quantified_impact = len(re.findall(r"\b\d+%|\b\d+x\b|\$\d+|\b\d+\+\b", text.lower()))
    verbs = sum(1 for item in experience_items if any(verb in normalize(item) for verb in ACTION_VERBS))
    job_years = max((int(match) for match in re.findall(r"(\d+)\+?\s*years?", job_description.lower())), default=0)
    resume_years = max((int(match) for match in re.findall(r"(\d+)\+?\s*years?", text.lower())), default=0)

    score = 45
    score += min(20, len(experience_items) * 4)
    score += min(15, years_mentions * 3)
    score += min(10, quantified_impact * 2)
    score += min(10, verbs * 2)
    if job_years:
        if resume_years >= job_years:
            score += 10
        elif resume_years > 0:
            score += max(2, int((resume_years / job_years) * 10))
    return min(100, score), []


def formatting_score(text: str) -> int:
    lines = split_lines(text)
    if not lines:
        return 0

    bullets = sum(1 for line in lines if line.startswith(("-", "•", "*")))
    headings = sum(1 for line in lines if normalize(line).rstrip(":") in SECTION_HEADINGS)
    has_sections = sum(1 for heading in SECTION_HEADINGS if re.search(rf"(?im)^\s*{re.escape(heading)}\s*:?\s*$", text))
    contact_score = 0
    if extract_email(text):
        contact_score += 10
    if extract_phone(text):
        contact_score += 10
    if extract_name(text):
        contact_score += 5

    score = 25 + min(20, bullets * 2) + min(20, headings * 4) + min(15, has_sections * 3) + contact_score
    if len(lines) >= 25:
        score += 5
    if len(lines) >= 50:
        score += 5
    return min(100, score)


def project_score(projects: List[str], text: str) -> Tuple[int, List[str]]:
    project_items = projects or extract_projects(text)
    if not project_items:
        return (35 if text.strip() else 0), ["Add 2-3 projects that show real outcomes and technologies used."]

    detail_bonus = 0
    for item in project_items:
        lower = normalize(item)
        if any(char.isdigit() for char in item):
            detail_bonus += 1
        if any(tech in lower for tech in ["react", "node", "python", "aws", "docker", "mongodb", "sql", "redis"]):
            detail_bonus += 1

    score = min(100, 45 + len(project_items) * 10 + min(10, detail_bonus * 2))
    return score, []


def build_suggestions(
    missing_keywords: List[str],
    missing_skills: List[str],
    experience_items: List[str],
    projects: List[str],
    formatting: int,
) -> List[str]:
    suggestions = []
    if missing_keywords:
        suggestions.append(f"Add role keywords such as {', '.join(missing_keywords[:5])}.")
    if missing_skills:
        suggestions.append(f"Close the skills gap for {', '.join(missing_skills[:5])}.")
    if len(experience_items) < 3:
        suggestions.append("Add more measurable achievements under experience and quantify business impact.")
    if len(projects) < 2:
        suggestions.append("Include 2-3 projects with technology, scale, and outcomes.")
    if formatting < 70:
        suggestions.append("Use clear section headings, bullets, and a stronger contact block.")
    if not suggestions:
        suggestions.append("Tailor the summary and achievements to the target role for a stronger ATS fit.")
    return suggestions[:8]


def build_strengths(
    keyword_score_value: int,
    skills_score_value: int,
    experience_score_value: int,
    formatting_score_value: int,
    project_score_value: int,
    skills: List[str],
    projects: List[str],
) -> List[str]:
    strengths = []
    if keyword_score_value >= 70:
        strengths.append("Strong keyword alignment with the target role.")
    if skills_score_value >= 70 and skills:
        strengths.append(f"Clear technical stack evidence across {', '.join(skills[:5])}.")
    if experience_score_value >= 70:
        strengths.append("Experience section shows good depth and impact.")
    if formatting_score_value >= 70:
        strengths.append("Resume structure is ATS friendly.")
    if project_score_value >= 70 and projects:
        strengths.append("Projects section adds relevant proof of work.")
    return strengths or ["The resume has a solid base, but it needs more job-specific detail."]


def build_weaknesses(
    keyword_score_value: int,
    skills_score_value: int,
    experience_score_value: int,
    formatting_score_value: int,
    project_score_value: int,
) -> List[str]:
    weaknesses = []
    if keyword_score_value < 65:
        weaknesses.append("Keyword alignment is below target.")
    if skills_score_value < 65:
        weaknesses.append("Skill coverage is incomplete for the role.")
    if experience_score_value < 65:
        weaknesses.append("Experience bullets need more measurable impact.")
    if formatting_score_value < 65:
        weaknesses.append("ATS formatting signals can be strengthened.")
    if project_score_value < 65:
        weaknesses.append("Projects do not yet demonstrate enough relevance.")
    return weaknesses


def build_job_match(
    resume_text: str,
    resume_skills: List[str],
    resume_experience: List[str],
    resume_certifications: List[str],
    job_description: str,
    required_skills: List[str],
    required_certifications: List[str],
    experience_text: str = "",
) -> Dict[str, Any]:
    job_skills, job_keywords, job_certs = extract_job_requirements(job_description, required_skills)
    job_certs = unique_preserve_order(required_certifications + job_certs)

    resume_skill_norm = {normalize(skill) for skill in resume_skills}
    resume_cert_norm = {normalize(cert) for cert in resume_certifications}

    matching_skills = [skill for skill in job_skills if normalize(skill) in resume_skill_norm]
    missing_skills = [skill for skill in job_skills if normalize(skill) not in resume_skill_norm]
    missing_certs = [cert for cert in job_certs if normalize(cert) not in resume_cert_norm]

    resume_experience_text = " ".join(resume_experience)
    experience_mentions = max(
        [int(match) for match in re.findall(r"(\d+)\+?\s*years?", f"{resume_experience_text} {experience_text}".lower())]
        or [0]
    )
    required_years = max([int(match) for match in re.findall(r"(\d+)\+?\s*years?", f"{job_description} {experience_text}".lower())] or [0])
    missing_experience = []
    if required_years and experience_mentions < required_years:
        missing_experience.append(f"{required_years}+ years of experience")

    keyword_hits = []
    missing_keywords = []
    resume_haystack = normalize(f"{resume_text} {' '.join(resume_skills)} {' '.join(resume_experience)}")
    for keyword in unique_preserve_order(job_keywords):
        if normalize(keyword) in resume_haystack:
            keyword_hits.append(keyword)
        else:
            missing_keywords.append(keyword)

    score = 35
    score += min(35, int((len(matching_skills) / max(1, len(job_skills))) * 35))
    score += min(20, int((len(keyword_hits) / max(1, len(job_keywords))) * 20))
    if required_years:
        if experience_mentions >= required_years:
            score += 10
        elif experience_mentions > 0:
            score += max(2, int((experience_mentions / required_years) * 10))
    if not missing_certs:
        score += 5
    score = min(100, score)

    return {
        "matchScore": score,
        "matchingSkills": matching_skills,
        "missingTechnologies": missing_skills,
        "missingCertifications": missing_certs[:10],
        "missingKeywords": unique_preserve_order(missing_keywords)[:15],
        "missingExperienceRequirements": missing_experience,
        "matchedKeywords": keyword_hits[:15],
        "summary": "Strong alignment" if score >= 80 else "Partial alignment with clear gaps to close",
    }


def semantic_match_score(resume_text: str, job_text: str) -> int:
    if not resume_text.strip() or not job_text.strip():
        return 0
    if EMBEDDER is not None:
        try:
            embeddings = EMBEDDER.encode([resume_text, job_text], normalize_embeddings=True)
            score = float(embeddings[0] @ embeddings[1].T)
            return int(round(max(0.0, min(1.0, score)) * 100))
        except Exception:
            logger.exception("Sentence transformer comparison failed")

    resume_tokens = set(normalize(resume_text).split())
    job_tokens = set(normalize(job_text).split())
    overlap = len(resume_tokens.intersection(job_tokens))
    return int(round((overlap / max(1, len(job_tokens))) * 100))


def calculate_ats_analysis(
    text: str,
    filename: str,
    job_description: str = "",
    target_role: str = "",
) -> Dict[str, Any]:
    extracted_name = extract_name(text)
    extracted_email = extract_email(text)
    extracted_phone = extract_phone(text)
    skills = extract_skills(text, job_description)
    experience = extract_experience(text)
    education = extract_education(text)
    projects = extract_projects(text)
    certifications = extract_certifications(text)
    job_skills, job_keywords, job_certs = extract_job_requirements(job_description)

    keyword_value, missing_keywords = keyword_score(text, job_description, skills, job_skills)
    skills_value, missing_skills = skills_score(skills, job_skills)
    experience_value, _ = experience_score(experience, text, job_description)
    formatting_value = formatting_score(text)
    project_value, project_weaknesses = project_score(projects, text)
    overall = int(
        round(
            (keyword_value * 0.35)
            + (skills_value * 0.25)
            + (experience_value * 0.20)
            + (formatting_value * 0.10)
            + (project_value * 0.10)
        )
    )

    job_match = build_job_match(
        resume_text=text,
        resume_skills=skills,
        resume_experience=experience,
        resume_certifications=certifications,
        job_description=job_description,
        required_skills=job_skills,
        required_certifications=job_certs,
        experience_text=target_role,
    )

    suggestions = build_suggestions(missing_keywords, missing_skills, experience, projects, formatting_value)
    if project_weaknesses:
        suggestions.extend(project_weaknesses)
    strengths = build_strengths(keyword_value, skills_value, experience_value, formatting_value, project_value, skills, projects)
    weaknesses = build_weaknesses(keyword_value, skills_value, experience_value, formatting_value, project_value)

    weak_sections = []
    if skills_value < 60:
        weak_sections.append("Skills")
    if experience_value < 60:
        weak_sections.append("Experience")
    if project_value < 60:
        weak_sections.append("Projects")
    if formatting_value < 60:
        weak_sections.append("Formatting")

    return {
        "analysisSource": "python",
        "rawText": text,
        "overallScore": overall,
        "keywordScore": keyword_value,
        "skillsScore": skills_value,
        "experienceScore": experience_value,
        "formattingScore": formatting_value,
        "projectScore": project_value,
        "strengths": unique_preserve_order(strengths),
        "weaknesses": unique_preserve_order(weaknesses),
        "missingKeywords": unique_preserve_order(missing_keywords + [keyword for keyword in job_keywords if normalize(keyword) not in normalize(text)])[:20],
        "suggestions": unique_preserve_order(suggestions),
        "weakSections": unique_preserve_order(weak_sections),
        "improvementAreas": unique_preserve_order(suggestions),
        "name": extracted_name,
        "email": extracted_email,
        "phone": extracted_phone,
        "skills": skills,
        "experience": experience,
        "education": education,
        "projects": projects,
        "certifications": certifications,
        "jobMatch": job_match,
        "debug": {
            "filename": filename,
            "textLength": len(text),
            "extractedSkills": skills[:25],
            "missingKeywords": missing_keywords[:25],
            "atsValues": {
                "overall": overall,
                "keywordScore": keyword_value,
                "skillsScore": skills_value,
                "experienceScore": experience_value,
                "formattingScore": formatting_value,
                "projectScore": project_value,
            },
        },
    }


def call_openai_json(system_prompt: str, user_prompt: str) -> Optional[Dict[str, Any]]:
    if not OPENAI_API_KEY:
        return None

    request_body = json.dumps(
        {
            "model": OPENAI_MODEL,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": 0.4,
            "response_format": {"type": "json_object"},
        }
    ).encode("utf-8")

    request = urllib.request.Request(
        "https://api.openai.com/v1/chat/completions",
        data=request_body,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {OPENAI_API_KEY}",
        },
        method="POST",
    )

    try:
        with urllib.request.urlopen(request, timeout=40) as response:
            payload = json.loads(response.read().decode("utf-8"))
            content = payload["choices"][0]["message"]["content"]
            return json.loads(content)
    except Exception:
        logger.exception("OpenAI request failed")
        return None


def build_optimized_resume(analysis: Dict[str, Any], job_description: str, target_role: str) -> Dict[str, Any]:
    strengths = analysis.get("strengths", [])
    missing_keywords = analysis.get("missingKeywords", [])
    skills = analysis.get("skills", [])
    projects = analysis.get("projects", [])
    experience = analysis.get("experience", [])
    target = target_role.strip() or "the target role"

    system_prompt = (
        "You are an ATS optimization engine. Return JSON only with keys: "
        "optimizedResume, improvements, rewrittenBullets, missingKeywords, predictedATSScore."
    )
    user_prompt = json.dumps(
        {
            "targetRole": target,
            "jobDescription": job_description,
            "resumeAnalysis": {
                "strengths": strengths,
                "missingKeywords": missing_keywords,
                "skills": skills,
                "experience": experience,
                "projects": projects,
                "summary": analysis.get("rawText", "")[:6000],
            },
            "goal": "Improve ATS score, keyword relevance, action verbs, achievements, and bullet point quality.",
        },
        ensure_ascii=False,
    )

    ai_payload = call_openai_json(system_prompt, user_prompt)
    if ai_payload:
        return {
            "analysisSource": "openai",
            "optimizedResume": str(ai_payload.get("optimizedResume", "")).strip(),
            "improvements": ai_payload.get("improvements", []),
            "rewrittenBullets": ai_payload.get("rewrittenBullets", []),
            "missingKeywords": unique_preserve_order(ai_payload.get("missingKeywords", []) or missing_keywords),
            "predictedATSScore": int(ai_payload.get("predictedATSScore", analysis.get("overallScore", 0))),
            "analysis": analysis,
        }

    bullet_lines = [line for line in experience if line.startswith(("-", "•", "*"))]
    rewritten_bullets = []
    for line in bullet_lines[:6]:
        cleaned = line.lstrip("-•* ").strip()
        if cleaned and not cleaned[0].isupper():
            cleaned = cleaned[:1].upper() + cleaned[1:]
        rewritten_bullets.append(f"Led impact-focused work: {cleaned}")

    if not rewritten_bullets:
        rewritten_bullets = [
            "Led cross-functional delivery with measurable business impact.",
            "Optimized workflows using the most relevant tools for the role.",
            "Improved stakeholder outcomes through automation and clear execution.",
        ]

    optimized_sections = [
        f"Target Role: {target}",
        "",
        "Professional Summary",
        f"Results-driven candidate aligned to {target} with experience in {', '.join(skills[:6]) or 'core technical delivery'}.",
        "",
        "Key Strengths",
        f"- {', '.join(strengths[:3]) or 'Strong ownership, communication, and delivery focus.'}",
        "",
        "Suggested Bullet Improvements",
    ]
    optimized_sections.extend(f"- {bullet}" for bullet in rewritten_bullets)
    if projects:
        optimized_sections.extend(["", "Projects", *[f"- {project}" for project in projects[:4]]])

    return {
        "analysisSource": "heuristic",
        "optimizedResume": "\n".join(optimized_sections).strip(),
        "improvements": [
            "Added a role-specific summary.",
            "Reframed experience bullets with action verbs.",
            "Highlighted the most relevant skills and projects.",
        ],
        "rewrittenBullets": rewritten_bullets,
        "missingKeywords": missing_keywords,
        "predictedATSScore": min(100, analysis.get("overallScore", 0) + 8),
        "analysis": analysis,
    }


def build_cover_letter(payload: CoverLetterRequest) -> Dict[str, Any]:
    target_role = payload.job_title.strip() or "the role"
    company = payload.company_name.strip() or "your company"
    skills = ", ".join(payload.skills[:6]) if payload.skills else "technical delivery, problem solving, and collaboration"
    experience = "; ".join(payload.experience[:3]) if payload.experience else "a track record of delivering meaningful results"

    system_prompt = (
        "You write professional ATS-friendly cover letters. Return JSON only with keys: coverLetter and tone."
    )
    user_prompt = json.dumps(
        {
            "resumeSummary": payload.resume_summary[:4000],
            "jobTitle": target_role,
            "companyName": company,
            "jobDescription": payload.job_description[:6000],
            "skills": payload.skills,
            "experience": payload.experience,
        },
        ensure_ascii=False,
    )

    ai_payload = call_openai_json(system_prompt, user_prompt)
    if ai_payload and ai_payload.get("coverLetter"):
        return {"coverLetter": ai_payload["coverLetter"], "tone": ai_payload.get("tone", "professional"), "analysisSource": "openai"}

    body = payload.job_description.strip()
    cover_letter = (
        f"Dear Hiring Manager,\n\n"
        f"I am excited to apply for {target_role} at {company}. "
        f"My background combines {skills} and {experience}, which aligns well with the responsibilities described. "
        f"{'The opportunity stood out because it emphasizes impact, ownership, and execution. ' if body else ''}"
        f"I would welcome the chance to discuss how I can contribute to your team.\n\n"
        f"Sincerely,\nCandidate"
    )
    return {"coverLetter": cover_letter, "tone": "professional", "analysisSource": "heuristic"}


def wrap_text_lines(text: str, width: int = 88) -> List[str]:
    lines: List[str] = []
    for paragraph in text.replace("\r", "").split("\n"):
        if not paragraph.strip():
            lines.append("")
            continue
        wrapped = textwrap.wrap(paragraph, width=width) or [""]
        lines.extend(wrapped)
    return lines


def escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def build_pdf_bytes(title: str, content: str) -> bytes:
    lines = [title.strip()] + [""] + wrap_text_lines(content, width=90)
    lines = lines[:ATS_EXPORT_MAX_LINES]
    if len(wrap_text_lines(content, width=90)) > ATS_EXPORT_MAX_LINES:
        lines.append("... truncated for export ...")

    # Minimal one-page PDF with built-in Helvetica font.
    content_lines = ["BT", "/F1 11 Tf", "72 770 Td"]
    first = True
    for line in lines:
        escaped = escape_pdf_text(line)
        if first:
            content_lines.append(f"({escaped}) Tj")
            first = False
        else:
            content_lines.append("T*")
            content_lines.append(f"({escaped}) Tj")
    content_lines.append("ET")
    stream = "\n".join(content_lines).encode("latin-1", errors="ignore")

    objects = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objects.append(
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
    )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    objects.append(
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream"
    )

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref_start = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(b"trailer\n")
    pdf.extend(f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode("ascii"))
    pdf.extend(b"startxref\n")
    pdf.extend(f"{xref_start}\n".encode("ascii"))
    pdf.extend(b"%%EOF")
    return bytes(pdf)


def build_docx_bytes(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in content.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
        else:
            doc.add_paragraph("")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_content(payload: ExportRequest) -> Tuple[bytes, str, str]:
    base_name = re.sub(r"[^a-z0-9._-]+", "-", payload.file_name.lower()).strip("-") or "export"
    if payload.format.lower() == "docx":
        return (
            build_docx_bytes(payload.title or "Resume", payload.content),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f'attachment; filename="{base_name}.docx"',
        )
    if payload.format.lower() == "pdf":
        return (
            build_pdf_bytes(payload.title or "Resume", payload.content),
            "application/pdf",
            f'attachment; filename="{base_name}.pdf"',
        )
    raise HTTPException(status_code=400, detail="Unsupported export format.")


@app.get("/health")
def health():
    return {"success": True, "message": "ATS service is healthy."}


@app.post("/analyze")
async def analyze(file: UploadFile = File(...), job_description: str = Form(""), target_role: str = Form("")):
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Empty resume file.")

    logger.info(
        "Analyze request filename=%s mimetype=%s bytes=%s jd_chars=%s target_role=%s",
        file.filename,
        file.content_type,
        len(file_bytes),
        len(job_description or ""),
        target_role,
    )

    text = extract_text(file, file_bytes)
    if not text.strip():
        logger.warning("No extractable text for filename=%s", file.filename)

    analysis = calculate_ats_analysis(text, file.filename or "", job_description=job_description, target_role=target_role)
    logger.info(
        "Analyze result filename=%s overall=%s keyword=%s skills=%s experience=%s formatting=%s project=%s missing=%s",
        file.filename,
        analysis["overallScore"],
        analysis["keywordScore"],
        analysis["skillsScore"],
        analysis["experienceScore"],
        analysis["formattingScore"],
        analysis["projectScore"],
        analysis["missingKeywords"][:8],
    )
    return analysis


@app.post("/optimize")
async def optimize(file: UploadFile = File(...), job_description: str = Form(""), target_role: str = Form("")):
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Empty resume file.")

    logger.info(
        "Optimize request filename=%s mimetype=%s bytes=%s jd_chars=%s target_role=%s",
        file.filename,
        file.content_type,
        len(file_bytes),
        len(job_description or ""),
        target_role,
    )

    text = extract_text(file, file_bytes)
    analysis = calculate_ats_analysis(text, file.filename or "", job_description=job_description, target_role=target_role)
    optimized = build_optimized_resume(analysis, job_description=job_description, target_role=target_role)
    logger.info(
        "Optimize result filename=%s source=%s predictedScore=%s missing=%s",
        file.filename,
        optimized["analysisSource"],
        optimized["predictedATSScore"],
        optimized["missingKeywords"][:8],
    )
    return optimized


@app.post("/generate-cover-letter")
def generate_cover_letter(payload: CoverLetterRequest):
    logger.info(
        "Cover letter request role=%s company=%s skills=%s",
        payload.job_title,
        payload.company_name,
        payload.skills[:6],
    )
    result = build_cover_letter(payload)
    logger.info("Cover letter result source=%s tone=%s", result.get("analysisSource"), result.get("tone"))
    return result


@app.post("/match-job")
def match_job(payload: MatchJobRequest):
    resume_text = payload.resume_text or " ".join(payload.resume_skills) or " ".join(payload.resume_experience)
    job_text = " ".join(
        [
            payload.job_title,
            payload.job_description,
            " ".join(payload.required_skills),
            " ".join(payload.required_certifications),
            payload.location,
            payload.experience,
        ]
    ).strip()
    semantic_score = semantic_match_score(resume_text, job_text)

    job_skills, job_keywords, job_certs = extract_job_requirements(payload.job_description, payload.required_skills)
    match_details = build_job_match(
        resume_text=resume_text,
        resume_skills=payload.resume_skills,
        resume_experience=payload.resume_experience,
        resume_certifications=[],
        job_description=payload.job_description,
        required_skills=job_skills,
        required_certifications=payload.required_certifications + job_certs,
        experience_text=payload.experience,
    )
    match_details["matchScore"] = max(match_details["matchScore"], semantic_score)
    match_details["matchingSkills"] = unique_preserve_order(match_details.get("matchingSkills", []))
    match_details["missingKeywords"] = unique_preserve_order(match_details.get("missingKeywords", []) + job_keywords)
    match_details["summary"] = "Strong fit" if match_details["matchScore"] >= 80 else "Needs targeted updates"
    logger.info(
        "Job match score=%s matching=%s missing=%s",
        match_details["matchScore"],
        match_details["matchingSkills"][:6],
        match_details["missingTechnologies"][:6],
    )
    return match_details


@app.post("/export")
def export_document(payload: ExportRequest):
    if not payload.content.strip():
        raise HTTPException(status_code=400, detail="Export content is empty.")

    buffer, content_type, disposition = export_content(payload)
    logger.info("Export generated format=%s file=%s bytes=%s", payload.format, payload.file_name, len(buffer))
    return Response(
        content=buffer,
        media_type=content_type,
        headers={"Content-Disposition": disposition},
    )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8000")))
